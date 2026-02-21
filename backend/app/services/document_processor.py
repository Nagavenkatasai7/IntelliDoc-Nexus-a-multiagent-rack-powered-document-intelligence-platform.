import hashlib
import io
import time
from pathlib import Path

from app.core.config import get_settings
from app.core.logging import get_logger

logger = get_logger(__name__)
settings = get_settings()


class DocumentProcessor:
    """Extracts text content from various document formats."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".png", ".jpg", ".jpeg", ".tiff"}

    @staticmethod
    def compute_content_hash(content: bytes) -> str:
        return hashlib.sha256(content).hexdigest()

    async def extract_text(self, file_content: bytes, filename: str) -> dict:
        """Extract text and metadata from a document.

        Returns dict with keys: text, pages, metadata, page_count
        """
        ext = Path(filename).suffix.lower()
        start = time.time()

        if ext == ".pdf":
            result = self._extract_pdf(file_content)
        elif ext in (".docx", ".doc"):
            result = self._extract_docx(file_content)
        elif ext == ".txt":
            result = self._extract_txt(file_content)
        elif ext in (".png", ".jpg", ".jpeg", ".tiff"):
            result = self._extract_image(file_content)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        # Sanitize extracted text: strip null bytes and other control chars
        # that PostgreSQL TEXT columns cannot store
        result = self._sanitize_result(result)

        elapsed_ms = int((time.time() - start) * 1000)
        result["processing_time_ms"] = elapsed_ms
        logger.info(
            "document_extracted",
            filename=filename,
            page_count=result.get("page_count"),
            chars=len(result.get("text", "")),
            elapsed_ms=elapsed_ms,
        )
        return result

    @staticmethod
    def _sanitize_text(text: str) -> str:
        """Remove null bytes and other problematic control characters."""
        # Remove null bytes (PostgreSQL cannot store \x00 in TEXT)
        text = text.replace("\x00", "")
        # Remove other non-printable control chars except newline/tab/carriage-return
        return "".join(
            ch for ch in text
            if ch in ("\n", "\r", "\t") or (ord(ch) >= 32) or (ord(ch) >= 128)
        )

    def _sanitize_result(self, result: dict) -> dict:
        """Sanitize all text content in the extraction result."""
        if "text" in result:
            result["text"] = self._sanitize_text(result["text"])
        if "pages" in result:
            for page in result["pages"]:
                if "content" in page:
                    page["content"] = self._sanitize_text(page["content"])
        if "metadata" in result:
            for key, value in result["metadata"].items():
                if isinstance(value, str):
                    result["metadata"][key] = self._sanitize_text(value)
        return result

    def _extract_pdf(self, content: bytes) -> dict:
        import pdfplumber

        pages = []
        full_text_parts = []
        metadata = {}

        with pdfplumber.open(io.BytesIO(content)) as pdf:
            metadata = {
                "title": pdf.metadata.get("Title", ""),
                "author": pdf.metadata.get("Author", ""),
                "creator": pdf.metadata.get("Creator", ""),
                "producer": pdf.metadata.get("Producer", ""),
            }

            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                tables = page.extract_tables() or []

                # Convert tables to text representation
                table_texts = []
                for table in tables:
                    rows = []
                    for row in table:
                        cells = [str(cell) if cell else "" for cell in row]
                        rows.append(" | ".join(cells))
                    table_texts.append("\n".join(rows))

                page_content = text
                if table_texts:
                    page_content += "\n\n[TABLE]\n" + "\n[/TABLE]\n\n[TABLE]\n".join(table_texts) + "\n[/TABLE]"

                pages.append({
                    "page_number": i + 1,
                    "content": page_content,
                    "has_tables": len(tables) > 0,
                })
                full_text_parts.append(page_content)

        return {
            "text": "\n\n".join(full_text_parts),
            "pages": pages,
            "metadata": metadata,
            "page_count": len(pages),
        }

    def _extract_docx(self, content: bytes) -> dict:
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(content))
        paragraphs = []
        full_text_parts = []

        metadata = {
            "title": doc.core_properties.title or "",
            "author": doc.core_properties.author or "",
            "created": str(doc.core_properties.created) if doc.core_properties.created else "",
        }

        current_section = None
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect headings as section boundaries
            if para.style and para.style.name.startswith("Heading"):
                current_section = text

            paragraphs.append({
                "content": text,
                "style": para.style.name if para.style else None,
                "section": current_section,
            })
            full_text_parts.append(text)

        # Extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_text = "[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]"
            full_text_parts.append(table_text)

        return {
            "text": "\n\n".join(full_text_parts),
            "pages": [{"page_number": 1, "content": "\n\n".join(full_text_parts)}],
            "metadata": metadata,
            "page_count": 1,
        }

    def _extract_txt(self, content: bytes) -> dict:
        text = content.decode("utf-8", errors="replace")
        return {
            "text": text,
            "pages": [{"page_number": 1, "content": text}],
            "metadata": {},
            "page_count": 1,
        }

    def _extract_image(self, content: bytes) -> dict:
        """Extract text from images using OCR."""
        try:
            import pytesseract
            from PIL import Image

            image = Image.open(io.BytesIO(content))
            text = pytesseract.image_to_string(image)
            return {
                "text": text,
                "pages": [{"page_number": 1, "content": text}],
                "metadata": {"ocr": True},
                "page_count": 1,
            }
        except ImportError:
            logger.warning("pytesseract not available, skipping OCR")
            return {
                "text": "[OCR not available - install pytesseract]",
                "pages": [],
                "metadata": {"ocr_error": "pytesseract not installed"},
                "page_count": 0,
            }
